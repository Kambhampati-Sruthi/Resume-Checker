import streamlit as st
import os
import json
import re
import sqlite3
import tempfile
import zipfile
from datetime import datetime
from typing import Dict, List, Tuple, Optional
import pandas as pd
import numpy as np
from dataclasses import dataclass, asdict
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import google.generativeai as genai
from io import BytesIO
import time

# Multiple PDF/Document processing libraries for better reliability
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False

# Page configuration
st.set_page_config(
    page_title="Resume Relevance Check System",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        text-align: center;
        color: #1f77b4;
        margin-bottom: 2rem;
    }
    
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
        margin: 0.5rem 0;
    }
    
    .score-excellent {
        color: #28a745;
        font-weight: bold;
    }
    
    .score-good {
        color: #ffc107;
        font-weight: bold;
    }
    
    .score-poor {
        color: #dc3545;
        font-weight: bold;
    }
    
    .shortlisted-card {
        background: linear-gradient(90deg, #28a745 0%, #20c997 100%);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
    
    .not-shortlisted-card {
        background: linear-gradient(90deg, #dc3545 0%, #fd7e14 100%);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
    
    .candidate-card {
        border: 1px solid #ddd;
        border-radius: 10px;
        padding: 1rem;
        margin: 0.5rem 0;
        background-color: #fafafa;
    }
    
    .stProgress > div > div > div > div {
        background-color: #1f77b4;
    }
</style>
""", unsafe_allow_html=True)

@dataclass
class JobDescription:
    title: str
    company: str
    skills: List[str]
    qualifications: List[str]
    experience: str
    description: str
    requirements: List[str]
    keywords: List[str]

@dataclass
class ResumeData:
    name: str
    email: str
    phone: str
    skills: List[str]
    experience: List[str]
    education: List[str]
    projects: List[str]
    certifications: List[str]
    raw_text: str
    filename: str

@dataclass
class ScreeningResult:
    candidate_name: str
    email: str
    phone: str
    filename: str
    overall_score: float
    hard_match_score: float
    semantic_score: float
    verdict: str
    feedback: str
    shortlisted: bool
    timestamp: str
    skill_matches: List[str]
    keyword_matches: List[str]
    detailed_breakdown: Dict
    shortlist_reason: str  # New field for shortlisting reason

class DocumentParser:
    """Enhanced document parser with multiple fallback options"""
    
    @staticmethod
    def extract_text_from_pdf_multiple_methods(file_bytes: bytes, filename: str = "document") -> str:
        """Try multiple PDF extraction methods"""
        text = ""
        methods_tried = []
        
        # Method 1: PyPDF2
        if PYPDF2_AVAILABLE and not text.strip():
            try:
                methods_tried.append("PyPDF2")
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                pass
        
        # Method 2: pdfplumber
        if PDFPLUMBER_AVAILABLE and not text.strip():
            try:
                methods_tried.append("pdfplumber")
                with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                pass
        
        # Method 3: PyMuPDF (fitz)
        if PYMUPDF_AVAILABLE and not text.strip():
            try:
                methods_tried.append("PyMuPDF")
                pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in pdf_doc:
                    text += page.get_text() + "\n"
                pdf_doc.close()
                if text.strip():
                    return text.strip()
            except Exception as e:
                pass
        
        # Method 4: textract (if available)
        if TEXTRACT_AVAILABLE and not text.strip():
            try:
                methods_tried.append("textract")
                # Save to temporary file for textract
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                    tmp_file.write(file_bytes)
                    tmp_file.flush()
                    text = textract.process(tmp_file.name).decode('utf-8')
                    os.unlink(tmp_file.name)
                if text.strip():
                    return text.strip()
            except Exception as e:
                pass
        
        return text.strip()

    @staticmethod
    def extract_text_from_docx_multiple_methods(file_bytes: bytes, filename: str = "document") -> str:
        """Try multiple DOCX extraction methods"""
        text = ""
        methods_tried = []
        
        # Method 1: python-docx
        if DOCX_AVAILABLE and not text.strip():
            try:
                methods_tried.append("python-docx")
                doc = docx.Document(BytesIO(file_bytes))
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                pass
        
        # Method 2: docx2txt
        if DOCX2TXT_AVAILABLE and not text.strip():
            try:
                methods_tried.append("docx2txt")
                # Save to temporary file for docx2txt
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    tmp_file.write(file_bytes)
                    tmp_file.flush()
                    text = docx2txt.process(tmp_file.name)
                    os.unlink(tmp_file.name)
                if text.strip():
                    return text.strip()
            except Exception as e:
                pass
        
        # Method 3: textract (if available)
        if TEXTRACT_AVAILABLE and not text.strip():
            try:
                methods_tried.append("textract")
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    tmp_file.write(file_bytes)
                    tmp_file.flush()
                    text = textract.process(tmp_file.name).decode('utf-8')
                    os.unlink(tmp_file.name)
                if text.strip():
                    return text.strip()
            except Exception as e:
                pass
        
        return text.strip()

    @classmethod
    def extract_text_from_uploaded_file(cls, uploaded_file) -> str:
        """Extract text from uploaded file with multiple fallback methods"""
        if not uploaded_file:
            return ""
        
        # Reset file pointer
        uploaded_file.seek(0)
        file_bytes = uploaded_file.read()
        
        if not file_bytes:
            return ""
        
        filename = uploaded_file.name
        file_type = uploaded_file.type
        
        if file_type == "application/pdf" or filename.lower().endswith('.pdf'):
            return cls.extract_text_from_pdf_multiple_methods(file_bytes, filename)
        elif (file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" 
              or filename.lower().endswith('.docx')):
            return cls.extract_text_from_docx_multiple_methods(file_bytes, filename)
        else:
            return ""

class GeminiParser:
    """Enhanced Gemini parser with better error handling and caching"""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.model = None
        if api_key:
            try:
                genai.configure(api_key=api_key)
                self.model = genai.GenerativeModel('gemini-pro')
            except Exception as e:
                st.error(f"Failed to initialize Gemini API: {e}")

    def parse_job_description(self, jd_text: str) -> JobDescription:
        """Parse job description with enhanced extraction"""
        if not jd_text.strip():
            return None
            
        if not self.model:
            return self._fallback_jd_parsing(jd_text)
            
        prompt = f"""
        Analyze this job description and extract structured information as JSON:

        Job Description:
        {jd_text}

        Extract ONLY valid JSON in this exact format:
        {{
            "title": "specific job title",
            "company": "company name",
            "skills": ["technical skill 1", "technical skill 2", ...],
            "qualifications": ["qualification 1", "qualification 2", ...],
            "experience": "experience requirements as string",
            "description": "brief job summary",
            "requirements": ["requirement 1", "requirement 2", ...],
            "keywords": ["important keyword 1", "important keyword 2", ...]
        }}

        Focus on technical skills, programming languages, tools, and frameworks.
        Return ONLY the JSON object, no other text.
        """

        try:
            response = self.model.generate_content(prompt)
            json_text = self._extract_json_from_response(response.text)
            jd_data = json.loads(json_text)
            
            jd = JobDescription(
                title=jd_data.get('title', 'Unknown Position'),
                company=jd_data.get('company', 'Unknown Company'),
                skills=jd_data.get('skills', []),
                qualifications=jd_data.get('qualifications', []),
                experience=jd_data.get('experience', 'Not specified'),
                description=jd_data.get('description', ''),
                requirements=jd_data.get('requirements', []),
                keywords=jd_data.get('keywords', [])
            )
            
            return jd
                
        except Exception as e:
            return self._fallback_jd_parsing(jd_text)

    def parse_resume(self, resume_text: str, filename: str = "") -> ResumeData:
        """Parse resume with enhanced extraction"""
        if not resume_text.strip():
            return None
            
        if not self.model:
            return self._fallback_resume_parsing(resume_text, filename)
            
        prompt = f"""
        Analyze this resume and extract structured information as JSON:

        Resume:
        {resume_text}

        Extract ONLY valid JSON in this exact format:
        {{
            "name": "full candidate name",
            "email": "email@domain.com",
            "phone": "phone number",
            "skills": ["skill 1", "skill 2", ...],
            "experience": ["job experience 1", "job experience 2", ...],
            "education": ["education 1", "education 2", ...],
            "projects": ["project 1", "project 2", ...],
            "certifications": ["cert 1", "cert 2", ...]
        }}

        Extract technical skills, programming languages, and work experiences.
        Return ONLY the JSON object, no other text.
        """

        try:
            response = self.model.generate_content(prompt)
            json_text = self._extract_json_from_response(response.text)
            resume_data = json.loads(json_text)
            
            resume = ResumeData(
                name=resume_data.get('name', 'Unknown Candidate'),
                email=resume_data.get('email', ''),
                phone=resume_data.get('phone', ''),
                skills=resume_data.get('skills', []),
                experience=resume_data.get('experience', []),
                education=resume_data.get('education', []),
                projects=resume_data.get('projects', []),
                certifications=resume_data.get('certifications', []),
                raw_text=resume_text,
                filename=filename
            )
            
            return resume
                
        except Exception as e:
            return self._fallback_resume_parsing(resume_text, filename)

    def _extract_json_from_response(self, response_text: str) -> str:
        """Extract JSON from response with better patterns"""
        # Remove markdown formatting
        response_text = re.sub(r'```json\s*', '', response_text)
        response_text = re.sub(r'```\s*', '', response_text)
        response_text = response_text.strip()
        
        # Find JSON object
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            return json_match.group(0)
        
        return response_text

    def _fallback_jd_parsing(self, jd_text: str) -> JobDescription:
        """Enhanced fallback parsing with better skill detection"""
        
        # Common technical skills to look for
        common_skills = [
            'python', 'java', 'javascript', 'react', 'angular', 'vue', 'node.js', 'nodejs',
            'sql', 'mysql', 'postgresql', 'mongodb', 'aws', 'azure', 'gcp', 'docker',
            'kubernetes', 'git', 'html', 'css', 'php', 'ruby', 'c++', 'c#', '.net',
            'django', 'flask', 'spring', 'express', 'tensorflow', 'pytorch', 'machine learning',
            'data science', 'analytics', 'tableau', 'power bi', 'excel', 'pandas', 'numpy'
        ]
        
        text_lower = jd_text.lower()
        found_skills = [skill for skill in common_skills if skill in text_lower]
        
        # Extract job title (usually in first few lines)
        title = self._extract_title(jd_text)
        company = self._extract_company(jd_text)
        
        # Extract requirements
        requirements = []
        req_patterns = [
            r'(?:requirements?|qualifications?)[\s:]+([^\n]+)',
            r'(?:must have|should have)[\s:]+([^\n]+)',
            r'(?:experience with)[\s:]+([^\n]+)'
        ]
        
        for pattern in req_patterns:
            matches = re.findall(pattern, jd_text, re.IGNORECASE)
            requirements.extend(matches)
        
        return JobDescription(
            title=title,
            company=company,
            skills=found_skills,
            qualifications=[],
            experience="Not specified",
            description=jd_text[:300] + "..." if len(jd_text) > 300 else jd_text,
            requirements=requirements,
            keywords=found_skills
        )

    def _fallback_resume_parsing(self, resume_text: str, filename: str = "") -> ResumeData:
        """Enhanced fallback parsing with better contact extraction"""
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, resume_text)
        email = email_match.group(0) if email_match else ""
        
        # Extract phone
        phone_patterns = [
            r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?\d{1,3}[-.\s]?\d{10}',
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}'
        ]
        
        phone = ""
        for pattern in phone_patterns:
            phone_match = re.search(pattern, resume_text)
            if phone_match:
                phone = phone_match.group(0)
                break
        
        # Extract name (usually first non-empty line)
        name = self._extract_name(resume_text)
        
        # Extract skills
        common_skills = [
            'python', 'java', 'javascript', 'react', 'angular', 'vue', 'node.js', 'nodejs',
            'sql', 'mysql', 'postgresql', 'mongodb', 'aws', 'azure', 'gcp', 'docker',
            'kubernetes', 'git', 'html', 'css', 'php', 'ruby', 'c++', 'c#', '.net'
        ]
        
        text_lower = resume_text.lower()
        found_skills = [skill for skill in common_skills if skill in text_lower]
        
        return ResumeData(
            name=name,
            email=email,
            phone=phone,
            skills=found_skills,
            experience=[],
            education=[],
            projects=[],
            certifications=[],
            raw_text=resume_text,
            filename=filename
        )

    def _extract_title(self, text: str) -> str:
        """Extract job title from text"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        job_keywords = ['developer', 'engineer', 'manager', 'analyst', 'specialist', 
                       'consultant', 'architect', 'lead', 'senior', 'junior']
        
        for line in lines[:10]:  # Check first 10 lines
            if any(keyword in line.lower() for keyword in job_keywords):
                return line.strip()
        
        # If no job keywords found, return first substantial line
        for line in lines[:5]:
            if len(line) > 10 and not any(char in line for char in '@()'):
                return line
        
        return "Unknown Position"

    def _extract_company(self, text: str) -> str:
        """Extract company name from text"""
        company_indicators = ['company:', 'organization:', 'employer:', 'about us', 'about the company']
        
        for indicator in company_indicators:
            pattern = f'{indicator}\\s*([^\\n]+)'
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                company = match.group(1).strip()
                if len(company) > 2:
                    return company
        
        return "Unknown Company"

    def _extract_name(self, text: str) -> str:
        """Extract candidate name from resume"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        for line in lines[:5]:  # Check first 5 lines
            # Skip lines with email or phone
            if '@' in line or any(char.isdigit() for char in line if char not in ' -+()'):
                continue
            
            # Check if line looks like a name (2-4 words, no special chars)
            words = line.split()
            if 2 <= len(words) <= 4 and all(word.isalpha() or word.endswith('.') for word in words):
                return line
        
        return "Unknown Candidate"

class RelevanceAnalyzer:
    """Enhanced relevance analyzer with detailed scoring"""
    
    def __init__(self):
        pass

    def hard_match_analysis(self, jd: JobDescription, resume: ResumeData) -> Tuple[float, Dict]:
        """Enhanced hard matching with comprehensive scoring"""
        
        match_details = {
            'skill_matches': [],
            'keyword_matches': [],
            'requirement_matches': [],
            'skills_score': 0,
            'keyword_score': 0,
            'requirements_score': 0,
            'total_jd_skills': len(jd.skills),
            'total_jd_keywords': len(jd.keywords),
            'total_requirements': len(jd.requirements)
        }

        # Normalize text for better matching
        resume_text_lower = resume.raw_text.lower()
        jd_skills_lower = [skill.lower().strip() for skill in jd.skills if skill.strip()]
        resume_skills_lower = [skill.lower().strip() for skill in resume.skills if skill.strip()]

        # Skills matching
        skill_matches = set()
        
        # Direct skill matches
        direct_matches = set(jd_skills_lower) & set(resume_skills_lower)
        skill_matches.update(direct_matches)
        
        # Text-based skill matches
        for skill in jd_skills_lower:
            if skill in resume_text_lower and skill not in skill_matches:
                skill_matches.add(skill)
        
        match_details['skill_matches'] = list(skill_matches)
        skills_score = len(skill_matches) / len(jd_skills_lower) if jd_skills_lower else 0
        match_details['skills_score'] = skills_score

        # Keywords matching
        jd_keywords_lower = [kw.lower().strip() for kw in jd.keywords if kw.strip()]
        keyword_matches = []
        
        for keyword in jd_keywords_lower:
            if keyword in resume_text_lower:
                keyword_matches.append(keyword)
        
        match_details['keyword_matches'] = keyword_matches
        keyword_score = len(keyword_matches) / len(jd_keywords_lower) if jd_keywords_lower else 0
        match_details['keyword_score'] = keyword_score

        # Requirements matching
        requirement_matches = []
        for req in jd.requirements:
            if req.lower().strip() in resume_text_lower:
                requirement_matches.append(req)
        
        match_details['requirement_matches'] = requirement_matches
        requirements_score = len(requirement_matches) / len(jd.requirements) if jd.requirements else 0
        match_details['requirements_score'] = requirements_score

        # Calculate weighted final score
        final_score = (skills_score * 50 + keyword_score * 30 + requirements_score * 20)
        
        return final_score, match_details

    def semantic_analysis(self, jd: JobDescription, resume: ResumeData) -> Tuple[float, Dict]:
        """Enhanced semantic analysis"""
        
        # Prepare text for analysis
        jd_text = f"{jd.title} {jd.description} {' '.join(jd.skills)} {' '.join(jd.requirements)}".lower()
        resume_combined = f"{' '.join(resume.skills)} {' '.join(resume.experience)} {' '.join(resume.projects)}".lower()
        
        # Word overlap analysis
        jd_words = set(re.findall(r'\b\w+\b', jd_text))
        resume_words = set(re.findall(r'\b\w+\b', resume_combined))
        
        # Remove common stop words
        stop_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
        jd_words -= stop_words
        resume_words -= stop_words
        
        common_words = jd_words & resume_words
        total_words = jd_words | resume_words
        
        similarity = len(common_words) / len(total_words) if total_words else 0
        semantic_score = similarity * 100
        
        return semantic_score, {
            'similarity': similarity,
            'common_words_count': len(common_words),
            'total_unique_words': len(total_words),
            'common_words': list(common_words)[:10]  # First 10 common words
        }

class DatabaseManager:
    """Enhanced database manager with better data handling"""
    
    def __init__(self):
        self.init_database()

    def init_database(self):
        """Initialize database in session state"""
        if 'screening_results' not in st.session_state:
            st.session_state.screening_results = []

    def store_result(self, result: ScreeningResult, jd: JobDescription):
        """Store result with enhanced data"""
        result_dict = {
            'candidate_name': result.candidate_name,
            'email': result.email,
            'phone': result.phone,
            'filename': result.filename,
            'overall_score': result.overall_score,
            'hard_match_score': result.hard_match_score,
            'semantic_score': result.semantic_score,
            'verdict': result.verdict,
            'feedback': result.feedback,
            'shortlisted': result.shortlisted,
            'timestamp': result.timestamp,
            'jd_title': jd.title,
            'jd_company': jd.company,
            'skill_matches': result.skill_matches,
            'keyword_matches': result.keyword_matches,
            'detailed_breakdown': result.detailed_breakdown,
            'shortlist_reason': result.shortlist_reason
        }
        st.session_state.screening_results.append(result_dict)

    def get_all_results(self) -> List[Dict]:
        """Get all screening results"""
        return st.session_state.get('screening_results', [])

    def get_shortlisted_candidates(self, limit: int = 50) -> List[Dict]:
        """Get shortlisted candidates"""
        all_results = self.get_all_results()
        shortlisted = [r for r in all_results if r['shortlisted']]
        return sorted(shortlisted, key=lambda x: x['overall_score'], reverse=True)[:limit]

    def clear_results(self):
        """Clear all results"""
        st.session_state.screening_results = []

class FeedbackGenerator:
    """Enhanced feedback generator"""
    
    def __init__(self, gemini_parser: GeminiParser):
        self.gemini_parser = gemini_parser

    def generate_feedback(self, jd: JobDescription, resume: ResumeData, 
                         screening_result: ScreeningResult) -> str:
        """Generate enhanced feedback"""
        if not self.gemini_parser.model:
            return self._generate_simple_feedback(jd, resume, screening_result)
            
        prompt = f"""
        Generate professional feedback for this candidate:

        Job: {jd.title}
        Required Skills: {', '.join(jd.skills[:10])}
        
        Candidate: {resume.name}
        Skills: {', '.join(resume.skills[:10])}
        
        Scores:
        - Overall: {screening_result.overall_score:.1f}%
        - Skills Match: {screening_result.hard_match_score:.1f}%
        - Shortlisted: {'Yes' if screening_result.shortlisted else 'No'}
        - Matched Skills: {', '.join(screening_result.skill_matches[:5])}

        Provide concise feedback (150 words max) covering:
        1. Key strengths
        2. Areas for improvement (specific skills to develop)
        3. Recommendation
        4. If not shortlisted, explain why and what they can improve

        Be constructive and professional.
        """

        try:
            response = self.gemini_parser.model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            return self._generate_simple_feedback(jd, resume, screening_result)

    def _generate_simple_feedback(self, jd: JobDescription, resume: ResumeData, 
                                 screening_result: ScreeningResult) -> str:
        """Generate analysis-based feedback using actual screening data"""
        feedback_parts = []
        
        # Extract detailed breakdown data
        hard_details = screening_result.detailed_breakdown.get('hard_details', {})
        semantic_details = screening_result.detailed_breakdown.get('semantic_details', {})
        
        # Overall performance analysis
        overall_score = screening_result.overall_score
        skills_score = hard_details.get('skills_score', 0) * 100
        semantic_score = screening_result.semantic_score
        
        # Score-based assessment with specific metrics
        if overall_score >= 75:
            feedback_parts.append(f"Excellent candidate with {overall_score:.1f}% overall match.")
        elif overall_score >= 60:
            feedback_parts.append(f"Strong candidate with {overall_score:.1f}% alignment.")
        elif overall_score >= 50:
            feedback_parts.append(f"Good potential with {overall_score:.1f}% match score.")
        elif overall_score >= 30:
            feedback_parts.append(f"Moderate fit at {overall_score:.1f}% - needs development.")
        else:
            feedback_parts.append(f"Limited alignment at {overall_score:.1f}% match.")
        
        # Detailed skills analysis
        skill_matches = screening_result.skill_matches
        total_jd_skills = hard_details.get('total_jd_skills', len(jd.skills))
        
        if skill_matches:
            skills_matched = len(skill_matches)
            skill_percentage = (skills_matched / total_jd_skills * 100) if total_jd_skills > 0 else 0
            
            if skills_matched >= 5:
                feedback_parts.append(f"Strong technical profile - matches {skills_matched}/{total_jd_skills} key skills ({skill_percentage:.0f}%).")
            elif skills_matched >= 3:
                feedback_parts.append(f"Good technical foundation - {skills_matched}/{total_jd_skills} skills matched ({skill_percentage:.0f}%).")
            elif skills_matched >= 1:
                feedback_parts.append(f"Some relevant skills - {skills_matched}/{total_jd_skills} matches found ({skill_percentage:.0f}%).")
            
            # List top matched skills
            top_skills = skill_matches[:4]  # Show top 4 skills
            feedback_parts.append(f"Strengths: {', '.join(top_skills)}.")
        else:
            feedback_parts.append("No direct skill matches found in resume content.")
        
        # Gap analysis with specific missing skills
        jd_skills_lower = set(skill.lower() for skill in jd.skills)
        matched_skills_lower = set(skill.lower() for skill in skill_matches)
        missing_skills = list(jd_skills_lower - matched_skills_lower)
        
        if missing_skills:
            priority_missing = missing_skills[:3]  # Top 3 missing skills
            if len(missing_skills) > 10:
                feedback_parts.append(f"Significant skill gaps identified - focus on core technologies like {', '.join(priority_missing)}.")
            elif len(missing_skills) > 5:
                feedback_parts.append(f"Several skill gaps - recommend developing {', '.join(priority_missing)}.")
            elif len(missing_skills) > 2:
                feedback_parts.append(f"Minor gaps in {', '.join(priority_missing)}.")
            else:
                feedback_parts.append(f"Close match - consider adding {', '.join(priority_missing)}.")
        
        # Semantic analysis insights
        common_words_count = semantic_details.get('common_words_count', 0)
        if common_words_count > 20:
            feedback_parts.append("Strong domain vocabulary alignment.")
        elif common_words_count > 10:
            feedback_parts.append("Good contextual understanding evident.")
        elif common_words_count > 5:
            feedback_parts.append("Some domain familiarity shown.")
        
        # Keyword analysis
        keyword_matches = screening_result.keyword_matches
        total_keywords = hard_details.get('total_jd_keywords', len(jd.keywords))
        if keyword_matches and total_keywords > 0:
            keyword_percentage = len(keyword_matches) / total_keywords * 100
            if keyword_percentage > 60:
                feedback_parts.append(f"Excellent keyword relevance ({keyword_percentage:.0f}%).")
            elif keyword_percentage > 30:
                feedback_parts.append(f"Good keyword alignment ({keyword_percentage:.0f}%).")
        
        # Shortlisting decision explanation
        if screening_result.shortlisted:
            reason = getattr(screening_result, 'shortlist_reason', 'Strong overall performance')
            feedback_parts.append(f"SHORTLISTED: {reason}.")
        else:
            reason = getattr(screening_result, 'shortlist_reason', 'Below minimum threshold')
            feedback_parts.append(f"Not shortlisted: {reason}.")
            
            # Specific improvement recommendations based on gaps
            if overall_score < 30:
                feedback_parts.append("Focus on developing core job-relevant skills and gaining more experience in this domain.")
            elif overall_score < 50:
                feedback_parts.append("Build stronger technical foundation and consider relevant certifications or projects.")
            else:
                feedback_parts.append("Close to shortlisting - strengthen key missing skills for future opportunities.")
        
        return " ".join(feedback_parts)

class ResumeScreeningSystem:
    """Enhanced screening system with batch processing"""
    
    def __init__(self, api_key: str):
        self.parser = GeminiParser(api_key)
        self.analyzer = RelevanceAnalyzer()
        self.db_manager = DatabaseManager()
        self.feedback_generator = FeedbackGenerator(self.parser)
        self.doc_parser = DocumentParser()

    def process_job_description(self, uploaded_file) -> JobDescription:
        """Process job description from uploaded file"""
        jd_text = self.doc_parser.extract_text_from_uploaded_file(uploaded_file)
        if jd_text:
            return self.parser.parse_job_description(jd_text)
        else:
            return None

    def process_resume(self, uploaded_file) -> ResumeData:
        """Process resume from uploaded file"""
        resume_text = self.doc_parser.extract_text_from_uploaded_file(uploaded_file)
        if resume_text:
            return self.parser.parse_resume(resume_text, uploaded_file.name)
        else:
            return None

    def screen_candidate(self, jd: JobDescription, resume: ResumeData) -> ScreeningResult:
        """Screen a single candidate with enhanced analysis"""
        
        # Hard Match Analysis
        hard_score, hard_details = self.analyzer.hard_match_analysis(jd, resume)
        
        # Semantic Analysis
        semantic_score, semantic_details = self.analyzer.semantic_analysis(jd, resume)
        
        # Calculate overall score with refined weighting
        overall_score = (hard_score * 0.75 + semantic_score * 0.25)
        
        # Enhanced shortlisting logic with reasons
        shortlist_reason = ""
        if overall_score >= 75:
            verdict = "Excellent Match"
            shortlisted = True
            shortlist_reason = "High overall score with strong skill alignment"
        elif overall_score >= 60:
            verdict = "Good Match"
            shortlisted = True
            shortlist_reason = "Good skill match and semantic alignment"
        elif overall_score >= 50 and len(hard_details.get('skill_matches', [])) >= 3:
            verdict = "Qualified Match"
            shortlisted = True
            shortlist_reason = "Multiple key skills matched despite moderate overall score"
        elif overall_score >= 45:
            verdict = "Fair Match"
            shortlisted = False
            shortlist_reason = "Below minimum threshold but shows potential"
        elif overall_score >= 30:
            verdict = "Poor Match"
            shortlisted = False
            shortlist_reason = "Limited skill alignment with job requirements"
        else:
            verdict = "No Match"
            shortlisted = False
            shortlist_reason = "Minimal alignment with required qualifications"

        # Create comprehensive screening result
        result = ScreeningResult(
            candidate_name=resume.name,
            email=resume.email,
            phone=resume.phone,
            filename=resume.filename,
            overall_score=overall_score,
            hard_match_score=hard_score,
            semantic_score=semantic_score,
            verdict=verdict,
            feedback="",
            shortlisted=shortlisted,
            timestamp=datetime.now().isoformat(),
            skill_matches=hard_details.get('skill_matches', []),
            keyword_matches=hard_details.get('keyword_matches', []),
            detailed_breakdown={
                'hard_details': hard_details,
                'semantic_details': semantic_details
            },
            shortlist_reason=shortlist_reason
        )

        # Generate feedback
        result.feedback = self.feedback_generator.generate_feedback(jd, resume, result)

        # Store in database
        self.db_manager.store_result(result, jd)

        return result

    def batch_process_resumes(self, jd: JobDescription, resume_files: List) -> List[ScreeningResult]:
        """Process multiple resumes with simplified progress tracking"""
        results = []
        
        # Single progress bar
        total_files = len(resume_files)
        progress_bar = st.progress(0)
        
        # Processing status
        status_placeholder = st.empty()
        
        for i, resume_file in enumerate(resume_files):
            # Update status
            status_placeholder.text(f"Processing {resume_file.name} ({i+1}/{total_files})")
            
            try:
                # Process resume
                resume = self.process_resume(resume_file)
                
                if resume:
                    # Screen candidate
                    result = self.screen_candidate(jd, resume)
                    results.append(result)
                    
                # Update progress
                progress = (i + 1) / total_files
                progress_bar.progress(progress)
                    
            except Exception as e:
                continue
        
        # Clear status and complete
        status_placeholder.empty()
        progress_bar.progress(1.0)
        st.success(f"Processed {len(results)}/{total_files} resumes")
        
        return results

def get_score_color(score: float) -> str:
    """Get color based on score"""
    if score >= 80:
        return "#28a745"  # Green
    elif score >= 65:
        return "#20c997"  # Teal
    elif score >= 50:
        return "#ffc107"  # Yellow
    elif score >= 35:
        return "#fd7e14"  # Orange
    else:
        return "#dc3545"  # Red

def create_score_gauge(score: float, title: str, unique_key: str = "") -> go.Figure:
    """Create an enhanced gauge chart for scores with unique key"""
    fig = go.Figure(go.Indicator(
        mode = "gauge+number+delta",
        value = score,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': title, 'font': {'size': 14, 'color': 'darkblue'}},
        delta = {'reference': 70, 'increasing': {'color': "green"}},
        gauge = {
            'axis': {'range': [None, 100], 'tickwidth': 1, 'tickcolor': "darkblue"},
            'bar': {'color': get_score_color(score)},
            'steps': [
                {'range': [0, 35], 'color': "#ffebee"},
                {'range': [35, 50], 'color': "#fff3e0"},
                {'range': [50, 65], 'color': "#fff8e1"},
                {'range': [65, 80], 'color': "#e8f5e8"},
                {'range': [80, 100], 'color': "#e3f2fd"}
            ],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': 90
            }
        }
    ))
    
    fig.update_layout(
        height=200,
        margin=dict(l=10, r=10, t=30, b=10),
        font={'color': "darkblue", 'family': "Arial"}
    )
    
    return fig

def main():
    # Navigation at top left - always visible
    st.sidebar.markdown("### Navigation")
    page = st.sidebar.selectbox(
        "Navigate to:",
        ["Home", "Analytics", "Results History", "Shortlisted", "Help"],
        key="main_navigation"
    )
    
    st.markdown('<h1 class="main-header">AI-Powered Resume Relevance Check System</h1>', unsafe_allow_html=True)
    
    # Get API key from Streamlit secrets
    api_key = ""
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
        st.sidebar.success("AI features enabled")
    except:
        st.sidebar.warning("AI features disabled - Add GEMINI_API_KEY to secrets.toml")
        st.sidebar.info("Basic text matching available")
    
    # Sidebar configuration
    with st.sidebar:
        st.header("Configuration")
        
        st.divider()
        
        # Available libraries info
        st.subheader("Available Libraries")
        libraries_status = {
            "PyPDF2": PYPDF2_AVAILABLE,
            "pdfplumber": PDFPLUMBER_AVAILABLE,  
            "PyMuPDF": PYMUPDF_AVAILABLE,
            "python-docx": DOCX_AVAILABLE,
            "docx2txt": DOCX2TXT_AVAILABLE,
            "textract": TEXTRACT_AVAILABLE
        }
        
        for lib, available in libraries_status.items():
            if available:
                st.success(f"✅ {lib}")
            else:
                st.error(f"❌ {lib}")
        
        installed_count = sum(libraries_status.values())
        st.info(f"{installed_count}/6 extraction libraries available")
        
        st.divider()
        
        # Clear data option
        if st.button("Clear All Data", type="secondary"):
            st.session_state.screening_results = []
            st.success("Data cleared!")
            st.rerun()

    # Initialize system
    screening_system = ResumeScreeningSystem(api_key)

    # Route to different pages
    if page == "Home":
        home_page(screening_system)
    elif page == "Analytics":
        analytics_page(screening_system)
    elif page == "Results History":
        results_history_page(screening_system)
    elif page == "Shortlisted":
        shortlisted_page(screening_system)
    elif page == "Help":
        help_page()

def home_page(screening_system):
    """Enhanced home page with batch processing"""
    st.header("Resume Screening & Analysis")
    
    # Create columns for file uploads
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("Job Description")
        jd_file = st.file_uploader(
            "Upload Job Description",
            type=['pdf', 'docx'],
            help="Upload the job description document (PDF or DOCX)",
            key="jd_upload"
        )
        
        if jd_file:
            with st.expander("Preview Job Description", expanded=False):
                jd_text = DocumentParser.extract_text_from_uploaded_file(jd_file)
                if jd_text:
                    preview_text = jd_text[:1000] + "..." if len(jd_text) > 1000 else jd_text
                    st.text_area("Content Preview", preview_text, height=200)
                    st.info(f"Total characters: {len(jd_text):,}")
    
    with col2:
        st.subheader("Resumes (Multiple Upload)")
        resume_files = st.file_uploader(
            "Upload Resume(s)",
            type=['pdf', 'docx'],
            accept_multiple_files=True,
            help="Upload multiple resume documents (PDF or DOCX)",
            key="resume_upload"
        )
        
        if resume_files:
            st.success(f"Uploaded {len(resume_files)} resume(s)")
            
            # Display uploaded files
            with st.expander("Uploaded Files", expanded=True):
                for i, file in enumerate(resume_files):
                    file_size = len(file.read()) if hasattr(file, 'read') else 0
                    file.seek(0)  # Reset file pointer
                    st.write(f"{i+1}. **{file.name}** ({file_size:,} bytes)")

    # Processing section
    if jd_file and resume_files:
        st.divider()
        
        # Processing options
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col1:
            min_score = st.slider("Minimum Score Filter", 0, 100, 0, 5)
        
        with col2:
            if st.button("Start Batch Screening", use_container_width=True, type="primary"):
                process_batch_screening(screening_system, jd_file, resume_files, min_score)
        
        with col3:
            auto_shortlist = st.checkbox("Auto-shortlist (≥60%)", value=True)

def process_batch_screening(screening_system, jd_file, resume_files, min_score_filter):
    """Process batch screening with enhanced UI"""
    
    st.divider()
    st.subheader("Processing Results")
    
    # Process job description
    jd = screening_system.process_job_description(jd_file)
    
    if not jd:
        st.error("Failed to process job description. Please try a different file.")
        return
    
    # Display job info
    st.success("Job Description Processed Successfully")
    with st.expander("Job Details", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Title:** {jd.title}")
            st.write(f"**Company:** {jd.company}")
            st.write(f"**Experience:** {jd.experience}")
        with col2:
            if jd.skills:
                st.write(f"**Required Skills:** {', '.join(jd.skills[:8])}")
                if len(jd.skills) > 8:
                    st.write(f"*...and {len(jd.skills) - 8} more*")
            if jd.keywords:
                st.write(f"**Keywords:** {', '.join(jd.keywords[:6])}")
    
    st.divider()
    
    # Process resumes in batch
    st.subheader("Processing Candidates")
    results = screening_system.batch_process_resumes(jd, resume_files)
    
    if not results:
        st.error("No results generated. Please check your resume files.")
        return
    
    # Filter results by minimum score
    filtered_results = [r for r in results if r.overall_score >= min_score_filter]
    
    if min_score_filter > 0:
        st.info(f"Showing {len(filtered_results)}/{len(results)} candidates with score ≥ {min_score_filter}%")
    
    # Display batch results
    display_batch_results(filtered_results if filtered_results else results)

def display_batch_results(results):
    """Display comprehensive batch screening results"""
    
    # Sort results by overall score
    results_sorted = sorted(results, key=lambda x: x.overall_score, reverse=True)
    
    # Summary metrics
    st.subheader("Batch Summary")
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        st.metric("Total Candidates", len(results))
    
    with col2:
        shortlisted_count = sum(1 for r in results if r.shortlisted)
        st.metric("Shortlisted", shortlisted_count)
    
    with col3:
        avg_score = np.mean([r.overall_score for r in results])
        st.metric("Avg Score", f"{avg_score:.1f}%")
    
    with col4:
        top_score = max([r.overall_score for r in results]) if results else 0
        st.metric("Top Score", f"{top_score:.1f}%")
    
    with col5:
        with_email = sum(1 for r in results if r.email)
        st.metric("With Email", with_email)
    
    # Score distribution chart
    st.subheader("Score Distribution")
    scores = [r.overall_score for r in results_sorted]
    names = [r.candidate_name[:20] + "..." if len(r.candidate_name) > 20 else r.candidate_name for r in results_sorted]
    
    fig = px.bar(
        x=names,
        y=scores,
        title="Candidate Screening Scores",
        color=scores,
        color_continuous_scale=['red', 'orange', 'yellow', 'lightgreen', 'green'],
        labels={'x': 'Candidates', 'y': 'Overall Score (%)'}
    )
    
    fig.update_layout(
        height=400,
        xaxis_tickangle=-45,
        showlegend=False,
        xaxis_title="Candidates",
        yaxis_title="Score (%)"
    )
    
    # Add score threshold lines
    fig.add_hline(y=60, line_dash="dash", line_color="green", 
                  annotation_text="Shortlist Threshold (60%)")
    fig.add_hline(y=50, line_dash="dash", line_color="orange", 
                  annotation_text="Fair Match (50%)")
    
    st.plotly_chart(fig, use_container_width=True, key="score_distribution_chart")
    
    # Quick actions
    st.subheader("Quick Actions")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("Generate Shortlist Report", use_container_width=True):
            generate_shortlist_report(results_sorted)
    
    with col2:
        if st.button("Export Email List", use_container_width=True):
            export_email_list(results_sorted)
    
    with col3:
        if st.button("Download Full Report", use_container_width=True):
            download_full_report(results_sorted)
    
    # Detailed candidate cards
    st.subheader("Detailed Candidate Analysis")
    
    # Tabs for different categories
    shortlisted_results = [r for r in results_sorted if r.shortlisted]
    not_shortlisted_results = [r for r in results_sorted if not r.shortlisted]
    
    tab1, tab2, tab3 = st.tabs([
        f"Shortlisted ({len(shortlisted_results)})",
        f"Not Shortlisted ({len(not_shortlisted_results)})",
        f"All Candidates ({len(results_sorted)})"
    ])
    
    with tab1:
        if shortlisted_results:
            for i, result in enumerate(shortlisted_results):
                display_candidate_card(result, is_shortlisted=True, card_index=f"shortlisted_{i}")
        else:
            st.info("No candidates were shortlisted based on current criteria.")
    
    with tab2:
        if not_shortlisted_results:
            for i, result in enumerate(not_shortlisted_results[:10]):  # Show first 10
                display_candidate_card(result, is_shortlisted=False, card_index=f"not_shortlisted_{i}")
            if len(not_shortlisted_results) > 10:
                st.info(f"Showing first 10 of {len(not_shortlisted_results)} candidates")
        else:
            st.info("All candidates were shortlisted!")
    
    with tab3:
        for i, result in enumerate(results_sorted):
            display_candidate_card(result, is_shortlisted=result.shortlisted, card_index=f"all_{i}")
            if i >= 19:  # Show first 20
                st.info(f"Showing first 20 of {len(results_sorted)} candidates")
                break

def display_candidate_card(result, is_shortlisted=True, card_index=""):
    """Display enhanced candidate card with unique keys"""
    
    with st.expander(
        f"{'✅' if is_shortlisted else '❌'} {result.candidate_name} - {result.verdict} ({result.overall_score:.1f}%)",
        expanded=False
    ):
        
        # Top row with key info
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.write("**Contact Information**")
            st.write(f"Email: {result.email or 'Not found'}")
            st.write(f"Phone: {result.phone or 'Not found'}")
            st.write(f"File: {result.filename}")
        
        with col2:
            st.write("**Scores**")
            st.write(f"Overall: **{result.overall_score:.1f}%**")
            st.write(f"Skills Match: **{result.hard_match_score:.1f}%**")
            st.write(f"Semantic: **{result.semantic_score:.1f}%**")
        
        with col3:
            st.write("**Status**")
            if result.shortlisted:
                st.success("SHORTLISTED")
                st.info(f"Reason: {getattr(result, 'shortlist_reason', 'High score')}")
            else:
                st.warning("NOT SHORTLISTED")
                st.info(f"Reason: {getattr(result, 'shortlist_reason', 'Below threshold')}")
            st.write(f"Verdict: {result.verdict}")
        
        # Score gauges with unique keys
        col1, col2, col3 = st.columns(3)
        
        with col1:
            fig1 = create_score_gauge(result.overall_score, "Overall", f"{card_index}_overall")
            st.plotly_chart(fig1, use_container_width=True, key=f"gauge_overall_{card_index}")
        
        with col2:
            fig2 = create_score_gauge(result.hard_match_score, "Skills Match", f"{card_index}_skills")
            st.plotly_chart(fig2, use_container_width=True, key=f"gauge_skills_{card_index}")
        
        with col3:
            fig3 = create_score_gauge(result.semantic_score, "Semantic", f"{card_index}_semantic")
            st.plotly_chart(fig3, use_container_width=True, key=f"gauge_semantic_{card_index}")
        
        # Skills analysis
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**Matched Skills**")
            if result.skill_matches:
                for skill in result.skill_matches:
                    st.success(f"✓ {skill}")
            else:
                st.info("No direct skill matches found")
            
            st.write("**Matched Keywords**")
            if result.keyword_matches:
                for keyword in result.keyword_matches:
                    st.info(f"• {keyword}")
            else:
                st.info("No keyword matches found")
        
        with col2:
            st.write("**AI Feedback**")
            st.write(result.feedback)
            
            # Action buttons for individual candidates
            button_col1, button_col2 = st.columns(2)
            with button_col1:
                if st.button(f"Email {result.candidate_name.split()[0]}", key=f"email_{card_index}"):
                    if result.email:
                        st.success(f"Email: {result.email}")
                        st.code(f"mailto:{result.email}?subject=Job Opportunity&body=Dear {result.candidate_name}, We would like to discuss...")
                    else:
                        st.warning("No email found for this candidate")
            
            with button_col2:
                if st.button("Detailed Report", key=f"report_{card_index}"):
                    show_detailed_candidate_report(result)

def show_detailed_candidate_report(result):
    """Show detailed candidate analysis"""
    st.write(f"## Detailed Report - {result.candidate_name}")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("### Contact Information")
        st.write(f"**Name:** {result.candidate_name}")
        st.write(f"**Email:** {result.email or 'Not provided'}")
        st.write(f"**Phone:** {result.phone or 'Not provided'}")
        st.write(f"**File:** {result.filename}")
        
        st.write("### Scoring Details")
        st.write(f"**Overall Score:** {result.overall_score:.1f}%")
        st.write(f"**Skills Match:** {result.hard_match_score:.1f}%")
        st.write(f"**Semantic Match:** {result.semantic_score:.1f}%")
        st.write(f"**Verdict:** {result.verdict}")
        
        if hasattr(result, 'shortlist_reason'):
            st.write(f"**Shortlist Reason:** {result.shortlist_reason}")
    
    with col2:
        st.write("### Skills Analysis")
        if result.skill_matches:
            st.write("**Matched Skills:**")
            for skill in result.skill_matches:
                st.write(f"• {skill}")
        
        if result.keyword_matches:
            st.write("**Matched Keywords:**")
            for keyword in result.keyword_matches:
                st.write(f"• {keyword}")
        
        st.write("### AI Feedback")
        st.write(result.feedback)

def generate_shortlist_report(results):
    """Generate downloadable shortlist report"""
    shortlisted = [r for r in results if r.shortlisted]
    
    if not shortlisted:
        st.warning("No candidates were shortlisted")
        return
    
    # Create DataFrame
    shortlist_data = []
    for result in shortlisted:
        shortlist_data.append({
            'Name': result.candidate_name,
            'Email': result.email,
            'Phone': result.phone,
            'Overall Score': f"{result.overall_score:.1f}%",
            'Skills Match': f"{result.hard_match_score:.1f}%",
            'Verdict': result.verdict,
            'Top Skills': ', '.join(result.skill_matches[:3]),
            'Shortlist Reason': getattr(result, 'shortlist_reason', 'High score'),
            'Filename': result.filename
        })
    
    df = pd.DataFrame(shortlist_data)
    
    st.success(f"Generated shortlist with {len(shortlisted)} candidates")
    st.dataframe(df, use_container_width=True)
    
    # Download button
    csv = df.to_csv(index=False)
    st.download_button(
        label="Download Shortlist CSV",
        data=csv,
        file_name=f"shortlisted_candidates_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
        mime="text/csv"
    )

def export_email_list(results):
    """Export email list of candidates"""
    candidates_with_email = [r for r in results if r.email]
    
    if not candidates_with_email:
        st.warning("No candidates with email addresses found")
        return
    
    # Create email list
    email_data = []
    for result in candidates_with_email:
        email_data.append({
            'Name': result.candidate_name,
            'Email': result.email,
            'Score': f"{result.overall_score:.1f}%",
            'Status': 'Shortlisted' if result.shortlisted else 'Not Shortlisted'
        })
    
    df = pd.DataFrame(email_data)
    
    st.success(f"Found {len(candidates_with_email)} candidates with email addresses")
    st.dataframe(df, use_container_width=True)
    
    # Email list for copying
    email_list = "; ".join([r.email for r in candidates_with_email])
    st.text_area("Email List (copy for BCC)", email_list, height=100)
    
    # Download button
    csv = df.to_csv(index=False)
    st.download_button(
        label="Download Email List CSV",
        data=csv,
        file_name=f"candidate_emails_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
        mime="text/csv"
    )

def download_full_report(results):
    """Generate comprehensive report for download"""
    # Create comprehensive DataFrame
    full_data = []
    for result in results:
        full_data.append({
            'Candidate Name': result.candidate_name,
            'Email': result.email,
            'Phone': result.phone,
            'Filename': result.filename,
            'Overall Score': result.overall_score,
            'Skills Match Score': result.hard_match_score,
            'Semantic Score': result.semantic_score,
            'Verdict': result.verdict,
            'Shortlisted': result.shortlisted,
            'Shortlist Reason': getattr(result, 'shortlist_reason', 'Standard evaluation'),
            'Matched Skills': ', '.join(result.skill_matches),
            'Matched Keywords': ', '.join(result.keyword_matches),
            'Feedback': result.feedback,
            'Timestamp': result.timestamp
        })
    
    df = pd.DataFrame(full_data)
    
    st.success(f"Generated full report with {len(results)} candidates")
    st.dataframe(df.head(), use_container_width=True)
    st.info("Showing first 5 rows. Download full report below.")
    
    # Download button
    csv = df.to_csv(index=False)
    st.download_button(
        label="Download Full Report CSV",
        data=csv,
        file_name=f"full_screening_report_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
        mime="text/csv"
    )

def shortlisted_page(screening_system):
    """Dedicated page for shortlisted candidates management"""
    
    st.header("Shortlisted Candidates")
    
    shortlisted_candidates = screening_system.db_manager.get_shortlisted_candidates()
    
    if not shortlisted_candidates:
        st.info("No shortlisted candidates yet. Run some screenings to see results here.")
        return
    
    # Summary metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Shortlisted", len(shortlisted_candidates))
    
    with col2:
        with_email = sum(1 for c in shortlisted_candidates if c.get('email'))
        st.metric("With Email", with_email)
    
    with col3:
        avg_score = np.mean([c['overall_score'] for c in shortlisted_candidates])
        st.metric("Average Score", f"{avg_score:.1f}%")
    
    with col4:
        top_score = max([c['overall_score'] for c in shortlisted_candidates])
        st.metric("Highest Score", f"{top_score:.1f}%")
    
    # Display candidates
    st.subheader("Shortlisted Candidates")
    
    for i, candidate in enumerate(shortlisted_candidates):
        display_shortlisted_candidate_card(candidate, i)

def display_shortlisted_candidate_card(candidate, index):
    """Display shortlisted candidate card with actions"""
    
    with st.container():
        st.markdown(f"""
        <div class="candidate-card">
            <h4>{candidate['candidate_name']} - {candidate['overall_score']:.1f}%</h4>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([2, 2, 1])
        
        with col1:
            st.write(f"**Email:** {candidate.get('email', 'Not available')}")
            st.write(f"**Phone:** {candidate.get('phone', 'Not available')}")
            st.write(f"**Verdict:** {candidate['verdict']}")
        
        with col2:
            st.write(f"**Skills Match:** {candidate['hard_match_score']:.1f}%")
            if candidate.get('skill_matches'):
                skills_display = ', '.join(candidate['skill_matches'][:3])
                st.write(f"**Top Skills:** {skills_display}")
            st.write(f"**Screened:** {candidate['timestamp'][:10]}")
        
        with col3:
            if candidate.get('email'):
                if st.button("Email", key=f"email_shortlisted_{index}"):
                    st.info(f"Email: {candidate['email']}")
            
            if st.button("Details", key=f"details_shortlisted_{index}"):
                show_shortlisted_details(candidate)

def show_shortlisted_details(candidate):
    """Show detailed shortlisted candidate information"""
    
    with st.expander(f"Detailed Report - {candidate['candidate_name']}", expanded=True):
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("### Contact Information")
            st.write(f"**Name:** {candidate['candidate_name']}")
            st.write(f"**Email:** {candidate.get('email', 'Not provided')}")
            st.write(f"**Phone:** {candidate.get('phone', 'Not provided')}")
            st.write(f"**File:** {candidate.get('filename', 'Unknown')}")
            
            st.write("### Scoring Details")
            st.write(f"**Overall Score:** {candidate['overall_score']:.1f}%")
            st.write(f"**Skills Match:** {candidate['hard_match_score']:.1f}%")
            st.write(f"**Semantic Match:** {candidate['semantic_score']:.1f}%")
            st.write(f"**Verdict:** {candidate['verdict']}")
            
            if candidate.get('shortlist_reason'):
                st.write(f"**Shortlist Reason:** {candidate['shortlist_reason']}")
        
        with col2:
            st.write("### Skills Analysis")
            if candidate.get('skill_matches'):
                st.write("**Matched Skills:**")
                for skill in candidate['skill_matches']:
                    st.write(f"• {skill}")
            
            if candidate.get('keyword_matches'):
                st.write("**Matched Keywords:**")
                for keyword in candidate['keyword_matches']:
                    st.write(f"• {keyword}")
            
            st.write("### AI Feedback")
            st.write(candidate.get('feedback', 'No feedback available'))

def analytics_page(screening_system):
    """Enhanced analytics page with comprehensive insights"""
    
    st.header("Analytics & Insights")
    
    all_results = screening_system.db_manager.get_all_results()
    
    if not all_results:
        st.info("No screening data available. Please run some screenings first.")
        return
    
    df = pd.DataFrame(all_results)
    
    # Key Performance Indicators
    st.subheader("Key Performance Indicators")
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        total_screened = len(df)
        st.metric("Total Screened", total_screened)
    
    with col2:
        shortlist_rate = (df['shortlisted'].sum() / len(df)) * 100 if len(df) > 0 else 0
        st.metric("Shortlist Rate", f"{shortlist_rate:.1f}%")
    
    with col3:
        avg_score = df['overall_score'].mean()
        st.metric("Avg Score", f"{avg_score:.1f}%")
    
    with col4:
        with_email_rate = (df['email'].str.len() > 0).sum() / len(df) * 100
        st.metric("Email Found", f"{with_email_rate:.1f}%")
    
    with col5:
        excellence_rate = (df['overall_score'] >= 80).sum() / len(df) * 100
        st.metric("Excellence Rate", f"{excellence_rate:.1f}%")
    
    # Score Distribution Analysis
    st.subheader("Score Distribution Analysis")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Score histogram
        fig_hist = px.histogram(
            df, x='overall_score', nbins=20,
            title="Overall Score Distribution",
            labels={'overall_score': 'Overall Score (%)', 'count': 'Number of Candidates'},
            color_discrete_sequence=['#1f77b4']
        )
        fig_hist.add_vline(x=60, line_dash="dash", line_color="green", 
                          annotation_text="Shortlist Threshold")
        st.plotly_chart(fig_hist, use_container_width=True, key="analytics_histogram")
    
    with col2:
        # Score categories pie chart
        score_categories = []
        for score in df['overall_score']:
            if score >= 80:
                score_categories.append('Excellent (80-100%)')
            elif score >= 60:
                score_categories.append('Good (60-79%)')
            elif score >= 45:
                score_categories.append('Fair (45-59%)')
            else:
                score_categories.append('Poor (0-44%)')
        
        category_counts = pd.Series(score_categories).value_counts()
        fig_pie = px.pie(
            values=category_counts.values,
            names=category_counts.index,
            title="Score Categories"
        )
        st.plotly_chart(fig_pie, use_container_width=True, key="analytics_pie_chart")

def results_history_page(screening_system):
    """Enhanced results history page"""
    
    st.header("Screening History")
    
    all_results = screening_system.db_manager.get_all_results()
    
    if not all_results:
        st.info("No screening history available.")
        return
    
    df = pd.DataFrame(all_results)
    df['timestamp'] = pd.to_datetime(df['timestamp']).dt.strftime('%Y-%m-%d %H:%M')
    
    # Display filtered results
    st.subheader(f"Results ({len(df)} candidates)")
    
    if len(df) > 0:
        # Display options
        display_cols = ['candidate_name', 'email', 'overall_score', 'verdict', 'shortlisted', 'feedback']
        
        display_df = df[display_cols].copy()
        
        # Format scores
        display_df['overall_score'] = display_df['overall_score'].apply(lambda x: f"{x:.1f}%")
        
        # Format shortlisted column
        display_df['shortlisted'] = display_df['shortlisted'].apply(lambda x: "✅" if x else "❌")
        
        st.dataframe(display_df, use_container_width=True)

def help_page():
    """Enhanced help page with troubleshooting"""
    
    st.header("Help & Documentation")
    
    # API Key Setup
    with st.expander("🔑 API Key Setup (Required for AI Features)", expanded=True):
        st.markdown("""
        ### Setting up Gemini API Key:
        
        1. **Get API Key**: Visit [Google AI Studio](https://ai.google.dev/) to get your free Gemini API key
        2. **Create secrets.toml file** in your project directory:
           ```
           # .streamlit/secrets.toml
           GEMINI_API_KEY = "your_api_key_here"
           ```
        3. **File structure should be**:
           ```
           your_project/
           ├── .streamlit/
           │   └── secrets.toml
           └── resume_screening.py
           ```
        4. **Restart the Streamlit app** after adding the secrets file
        
        ### Alternative for Streamlit Cloud:
        - Go to your app settings
        - Add secret: `GEMINI_API_KEY = "your_key"`
        - Deploy the updated app
        
        **Note**: Without API key, the system uses basic text matching instead of AI analysis.
        """)
    
    # Quick Start Guide
    with st.expander("🚀 Quick Start Guide", expanded=True):
        st.markdown("""
        ### Getting Started in 4 Easy Steps:
        
        1. **Setup API**: Add Gemini API key to secrets.toml (see above)
        2. **Upload JD**: Upload your job description (PDF or DOCX format)
        3. **Upload Resumes**: Select multiple resume files to screen at once
        4. **Start Screening**: Click the "Start Batch Screening" button and wait for results!
        """)
    
    # Shortlisting Logic
    with st.expander("🎯 Enhanced Shortlisting Logic", expanded=True):
        st.markdown("""
        ### How Candidates are Shortlisted:
        
        **Automatic Shortlisting Criteria:**
        - **75%+ Overall Score**: Excellent match - automatically shortlisted
        - **60-74% Overall Score**: Good match - automatically shortlisted  
        - **50-59% with 3+ Key Skills**: Qualified match - shortlisted despite moderate score
        - **45-49%**: Fair match - not shortlisted but shows potential
        - **Below 45%**: Poor/No match - not shortlisted
        
        **Enhanced Features:**
        - Detailed feedback based on actual analysis data (not templates)
        - Specific skill gap analysis with missing skills identified
        - Keyword relevance percentages and domain vocabulary assessment
        - Specific reasons provided for all shortlisting decisions
        - Data-driven recommendations for improvement
        """)
    
    # Installation Guide
    with st.expander("💻 Installation & Setup", expanded=False):
        st.markdown("""
        ### Required packages:
        ```bash
        pip install streamlit pandas plotly google-generativeai PyPDF2 python-docx numpy
        ```
        
        ### Optional (for better PDF handling):
        ```bash
        pip install pdfplumber PyMuPDF docx2txt textract
        ```
        
        ### Run the application:
        ```bash
        streamlit run resume_screening.py
        ```
        """)
    
    # Feedback System
    with st.expander("📊 AI Feedback System", expanded=False):
        st.markdown("""
        ### How Feedback is Generated:
        
        **With AI (Gemini API):**
        - Contextual analysis of candidate profile vs job requirements
        - Personalized recommendations based on specific gaps
        - Professional tone with constructive suggestions
        
        **Without AI (Fallback):**
        - Analysis-driven feedback using actual screening data
        - Specific skill gap identification with percentages
        - Keyword relevance analysis and domain vocabulary assessment
        - Data-based improvement recommendations
        - Detailed shortlisting explanations with specific reasons
        
        **Feedback includes:**
        - Overall performance score and breakdown
        - Specific matched skills and percentages
        - Priority missing skills for development
        - Semantic and keyword alignment analysis
        - Clear shortlisting decision reasoning
        """)
    
    st.subheader("Need Help?")
    col1, col2 = st.columns(2)
    
    with col1:
        st.info("""
        **Common Issues:**
        1. Ensure secrets.toml is in .streamlit/ folder
        2. Check file formats (PDF/DOCX only)
        3. Restart app after adding API key
        4. Verify files aren't password-protected
        """)
    
    with col2:
        st.success("""
        **Performance Tips:**
        1. Process 10-20 resumes at once
        2. Use well-formatted documents
        3. Include specific skills in job descriptions
        4. Review AI feedback for insights
        """)


if __name__ == "__main__":
    # Initialize session state
    if 'screening_results' not in st.session_state:
        st.session_state.screening_results = []
    
    # Run the main application
    main()
